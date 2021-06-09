import os
from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_LINE_SPACING
print(" ")
print(" ")
print(" ")
print(" ")
print("==============================================")
print("=== Acepta Cargo / Solicita Anticipo - PJN ===")
print("==============================================")
print(" ")
ans=True
while ans:
    print("Selecione un departamento")
    print ("""
    1. PJN
    2. San Isidro
    3. Pilar
    4. San Martin
    5. San Miguel
    """)
    dpto=input("Departamento: ") 
    if dpto=="1": 
      constituido=("Lavalle 1392, entrepiso, Casillero 652, C.A.B.A, provincia de Buenos Aires")
      ans=False 
    elif dpto=="2":
      constituido=("Rivadavia 220, Casillero 4054, San Isidro, provincia de Buenos Aires")
      ans=False  
    elif dpto=="3":
      constituido=("Lorenzo Lopez 355, piso 4 dpto A, Pilar, provincia de Buenos Aires")
      ans=False 
    elif dpto=="4":
      constituido=("Mitre 3586, piso 1° oficina “B”, casillero 4054, San Martín, provincia de Buenos Aires")
      ans=False  
    elif dpto=="5":
      constituido=("Av. Ricardo Balbin 1350, 1er piso, oficina 10, casillero 4054, San Miguel, provincia de Buenos Aires")
      ans=False             
    elif dpto !="":
      print("\n Opciòn invalida ingrese nuevamente") 

fechaNotificacion = input('Fecha de la notifiación: ')
dependencia = input('Dependencia: ')
numeroExpediente = input('Número de Expediente: ')
nombreExpediente = input('Nombre del Expediente: ')
nombreExpedienteFormateado=nombreExpediente.replace('/','--');

try:
    
    os.makedirs(nombreExpedienteFormateado+'/Escritos', exist_ok=True)
    os.makedirs(nombreExpedienteFormateado+'/Documental', exist_ok=True)

except OSError:
    print("La creación del directorio %s falló" % nombreExpedienteFormateado)

else:

	# Creación del documento de Aceptacion del Cargo
	documentAcepta = Document()

	# Fuente y tamaño
	paragraph = documentAcepta.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 2 
	run = paragraph.add_run(dependencia)
	run.bold = True

	paragraph = documentAcepta.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 2 # for left, 1 for center, 2 right, 3 justify ....
	run = paragraph.add_run('Expediente Nº '+numeroExpediente)
	run.bold = True

	paragraph = documentAcepta.add_paragraph()
	paragraph = documentAcepta.add_paragraph()

	paragraph = documentAcepta.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	run = paragraph.add_run('ACEPTA CARGO')
	run.underline = True
	run.bold = True

	paragraph = documentAcepta.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	run = paragraph.add_run('SEÑOR JUEZ')
	run.bold = True

	paragraph = documentAcepta.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	run = paragraph.add_run('	Gabriel Adrián Ascazuri')
	run.bold = True
	run = paragraph.add_run(', perito Ingeniero de Sistemas, D.N.I 26439965, C.U.I.T 20-26439965-4, Domicilio electrónico: 20264399654, Monotributista, con domicilio constituido en la calle '+constituido+', email: gapericias@gmail.com, Tel 1141980909, designado en los autos caratulados: ')
	run = paragraph.add_run('"'+nombreExpediente+'"')
	run.bold = True
	run = paragraph.add_run(', a V.S. respetuosamente digo:')
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	paragraph = documentAcepta.add_paragraph('	Que conforme fuese notificado en la cédula electrónica del día '+fechaNotificacion+' de la desinsaculación de este perito, con fecha de  sorteo del mismo día, vengo a aceptar el cargo para el que fui designado en autos, prometiendo desempeñar fielmente el mismo, acorde lo establece la normativa vigente.')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	paragraph = documentAcepta.add_paragraph('	Proveer de conformidad, que')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	paragraph = documentAcepta.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 1 # for left, 1 for center, 2 right, 3 justify ....
	run = paragraph.add_run('SERÁ JUSTICIA')
	run.bold = True

	#Ajusta Fuente y Tamaño
	style = documentAcepta.styles['Normal']
	font = style.font
	font.name = 'Verdana'
	font.size = Pt(12)

	#Ajusta Margenes
	sections = documentAcepta.sections
	for section in sections:
	    section.top_margin = Cm(5)
	    section.bottom_margin = Cm(2.5)
	    section.left_margin = Cm(5)
	    section.right_margin = Cm(2.5)

	documentAcepta.save(nombreExpedienteFormateado+'/Escritos/AceptaCargo.docx') 

	# FIN Creación del documento de Aceptacion del Cargo

	# ==================================================
	
	# Creación del documento de Solicita Anticipo
	documentSolicita = Document()

	# Fuente y tamaño
	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE	
	paragraph.alignment = 2 
	run = paragraph.add_run(dependencia)
	run.bold = True

	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE	
	paragraph.alignment = 2 # for left, 1 for center, 2 right, 3 justify ....
	run = paragraph.add_run('Expediente Nº '+numeroExpediente)
	run.bold = True

	paragraph = documentSolicita.add_paragraph()
	paragraph = documentSolicita.add_paragraph()

	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE	
	run = paragraph.add_run('SOLICITA ANTICIPO DE GASTOS – SE SUSPENDA - SOLICITA PRESTAMO EXPEDIENTE – AUTORIZA')
	run.underline = True
	run.bold = True

	paragraph = documentSolicita.add_paragraph()

	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE	
	run = paragraph.add_run('SEÑOR JUEZ')
	run.bold = True

	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE	
	run = paragraph.add_run('	Gabriel Adrián Ascazuri')
	run.bold = True
	run = paragraph.add_run(', perito Ingeniero de Sistemas, D.N.I 26439965, C.U.I.T 20-26439965-4, Domicilio electrónico: 20264399654, Monotributista, con domicilio constituido en la calle '+constituido+', email: gapericias@gmail.com, Tel 1141980909, designado en los autos caratulados: ')
	run = paragraph.add_run('"'+nombreExpediente+'"')
	run.bold = True
	run = paragraph.add_run(', a V.S. respetuosamente digo:')
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE	
	run = paragraph.add_run('	I.-')
	run.bold = True

	paragraph = documentSolicita.add_paragraph('	Que en atención a la aceptación del cargo de autos, vengo a solicitar la suma de PESOS CINCO MIL ($5.000,00.-) en concepto de anticipo de gastos.')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 
	
	paragraph = documentSolicita.add_paragraph('	Que dicho monto será utilizado para solventar los gastos de la pericia encomendada.')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph = documentSolicita.add_paragraph('	Que a dichos efectos denuncio cuenta bancaria a mi nombre, caja de Ahorro en pesos $000000120205112818 del Banco de la Cuidad de Buenos Aires, C.B.U  0290012410000051128189.')
	paragraph.alignment = 3


	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	run = paragraph.add_run('	II.-')
	run.bold = True
	paragraph = documentSolicita.add_paragraph('	Que solicito se suspendan los plazos para cumplir con la tarea encomendada, hasta tanto sea cumplido el depósito solicitado para solventar el monto de anticipo de gastos.')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	
	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	run = paragraph.add_run('	III.-')
	run.bold = True
	paragraph = documentSolicita.add_paragraph('	Que solicito se me otorgue en préstamo el expediente por el término de diez días.')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	
	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	run = paragraph.add_run('	IV.-')
	run.bold = True
	paragraph = documentSolicita.add_paragraph('	Que vengo a autorizar a compulsar a las presentes actuaciones a las Dra. JORGELINA MONICA REYES, AGUSTINA DIEZ, LAURA RODRIGUEZ, EMILIANO JUORIO y/o GONZALO DIEZ.')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	paragraph = documentSolicita.add_paragraph()

	paragraph = documentSolicita.add_paragraph('	Proveer de conformidad, que')
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....

	paragraph = documentSolicita.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
	paragraph.alignment = 1 # for left, 1 for center, 2 right, 3 justify ....
	run = paragraph.add_run('SERÁ JUSTICIA')
	run.bold = True

	#Ajusta Fuente y Tamaño
	style = documentSolicita.styles['Normal']
	font = style.font
	font.name = 'Verdana'
	font.size = Pt(12)

	#Ajusta Margenes
	sections = documentSolicita.sections
	for section in sections:
	    section.top_margin = Cm(5)
	    section.bottom_margin = Cm(2.5)
	    section.left_margin = Cm(5)
	    section.right_margin = Cm(2.5)

	documentSolicita.save(nombreExpedienteFormateado+'/Escritos/SolicitaAnticipo.docx') 

	# FIN Creación del documento de Solicita Anticipo
